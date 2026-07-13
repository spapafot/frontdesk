import asyncio
import io
import os
import re
import subprocess
import tempfile
from datetime import datetime, timezone

from sqlalchemy.ext.asyncio import AsyncSession

from app.core.config import settings
from app.models.knowledge import KnowledgeDocument
from app.repositories.knowledge_repository import KnowledgeRepository
from app.services.embeddings import embed_passage

SUPPORTED_EXTENSIONS = {".txt", ".pdf", ".doc", ".docx", ".xls", ".xlsx"}


class UnsupportedFileType(ValueError):
    pass


class ExtractionError(ValueError):
    pass


def _extension(filename: str) -> str:
    return os.path.splitext(filename)[1].lower()


def _extract_txt(data: bytes) -> str:
    return data.decode("utf-8", errors="ignore")


def _extract_pdf(data: bytes) -> str:
    from pypdf import PdfReader
    from pypdf.errors import FileNotDecryptedError

    reader = PdfReader(io.BytesIO(data))
    try:
        return "\n".join(page.extract_text() or "" for page in reader.pages)
    except FileNotDecryptedError as exc:
        raise ExtractionError(
            "This PDF is password-protected. Please upload an unlocked copy."
        ) from exc


def _extract_docx(data: bytes) -> str:
    from docx import Document

    document = Document(io.BytesIO(data))
    parts = [p.text for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def _extract_xlsx(data: bytes) -> str:
    from openpyxl import load_workbook

    workbook = load_workbook(io.BytesIO(data), read_only=True, data_only=True)
    lines: list[str] = []
    for sheet in workbook.worksheets:
        lines.append(f"# {sheet.title}")
        for row in sheet.iter_rows(values_only=True):
            cells = [str(c) for c in row if c is not None]
            if cells:
                lines.append(" | ".join(cells))
    return "\n".join(lines)


def _extract_xls(data: bytes) -> str:
    import xlrd

    workbook = xlrd.open_workbook(file_contents=data)
    lines: list[str] = []
    for sheet in workbook.sheets():
        lines.append(f"# {sheet.name}")
        for row_idx in range(sheet.nrows):
            cells = [str(c) for c in sheet.row_values(row_idx) if str(c).strip()]
            if cells:
                lines.append(" | ".join(cells))
    return "\n".join(lines)


def _extract_doc(data: bytes) -> str:
    # Legacy binary .doc -> use the antiword CLI installed in the image.
    with tempfile.NamedTemporaryFile(suffix=".doc", delete=False) as tmp:
        tmp.write(data)
        tmp_path = tmp.name
    try:
        result = subprocess.run(
            ["antiword", tmp_path],
            capture_output=True,
            timeout=60,
        )
        if result.returncode != 0:
            raise ExtractionError(
                "Could not read this .doc file. Please convert it to .docx or PDF."
            )
        return result.stdout.decode("utf-8", errors="ignore")
    finally:
        os.unlink(tmp_path)


_EXTRACTORS = {
    ".txt": _extract_txt,
    ".pdf": _extract_pdf,
    ".docx": _extract_docx,
    ".xlsx": _extract_xlsx,
    ".xls": _extract_xls,
    ".doc": _extract_doc,
}


def normalize_text(text: str) -> str:
    """Clean up extracted text: collapse runs of spaces/tabs, trim per-line
    whitespace, and limit consecutive blank lines (helps PDFs that emit double
    spaces between words, which otherwise degrade embeddings and matching)."""
    lines = [re.sub(r"[ \t]+", " ", line).strip() for line in text.splitlines()]
    cleaned = "\n".join(lines)
    cleaned = re.sub(r"\n{3,}", "\n\n", cleaned)
    return cleaned.strip()


def extract_text(filename: str, data: bytes) -> str:
    ext = _extension(filename)
    extractor = _EXTRACTORS.get(ext)
    if extractor is None:
        raise UnsupportedFileType(f"Unsupported file type: {ext or 'unknown'}")
    try:
        raw = extractor(data)
    except (UnsupportedFileType, ExtractionError):
        raise
    except Exception as exc:
        # Extraction is deterministic on the file bytes, so a parser failure
        # (corrupt/encrypted PDF, malformed docx, ...) will never succeed on
        # retry. Surface it as a terminal ExtractionError so the document is
        # marked "failed" on the first SQS receive instead of churning through
        # the retry/DLQ window stuck at "processing".
        raise ExtractionError(
            f"Could not read this {ext.lstrip('.') or 'file'} file: {exc}"
        ) from exc
    text = normalize_text(raw)
    if not text:
        raise ExtractionError("No readable text could be extracted from this file.")
    return text


def _is_useful_chunk(chunk: str) -> bool:
    """Drop low-signal chunks that pollute retrieval: near-empty fragments and
    lines dominated by punctuation/whitespace (e.g. PDF tables-of-contents with
    dotted leaders and page numbers)."""
    stripped = chunk.strip()
    if len(stripped) < 15:
        return False
    # Fraction of "real" characters (letters/digits) vs. dots, dashes, spaces.
    meaningful = sum(1 for c in stripped if c.isalnum())
    if meaningful / len(stripped) < 0.45:
        return False
    return True


def chunk_text(
    text: str, size: int | None = None, overlap: int | None = None
) -> list[str]:
    size = size if size is not None else settings.chunk_size
    overlap = overlap if overlap is not None else settings.chunk_overlap
    paragraphs = [p.strip() for p in text.split("\n") if p.strip()]
    chunks: list[str] = []
    current = ""
    for paragraph in paragraphs:
        if current and len(current) + len(paragraph) + 1 > size:
            chunks.append(current)
            tail = current[-overlap:] if overlap else ""
            current = f"{tail} {paragraph}".strip()
        else:
            current = f"{current}\n{paragraph}".strip() if current else paragraph
    if current:
        chunks.append(current)

    # Hard-split any oversized chunk (e.g. a single very long paragraph).
    final: list[str] = []
    for chunk in chunks:
        if len(chunk) <= size:
            final.append(chunk)
            continue
        start = 0
        while start < len(chunk):
            final.append(chunk[start : start + size])
            start += size - overlap

    return [chunk for chunk in final if _is_useful_chunk(chunk)]


async def ingest_document(
    session: AsyncSession, profile_id: int, filename: str, data: bytes
) -> tuple[KnowledgeDocument, int]:
    ext = _extension(filename)
    text = await asyncio.to_thread(extract_text, filename, data)
    chunks = chunk_text(text)

    repo = KnowledgeRepository(session)
    document = await repo.create_document(
        profile_id=profile_id,
        title=filename,
        type=ext.lstrip("."),
        content=text,
        is_active=True,
    )

    for chunk in chunks:
        embedding = await embed_passage(chunk)
        await repo.add_chunk(
            profile_id=profile_id,
            document_id=document.id,
            content=chunk,
            embedding=embedding,
            meta={"title": filename},
        )

    await session.commit()
    return document, len(chunks)


async def process_existing_document(
    session: AsyncSession,
    document: KnowledgeDocument,
    filename: str,
    data: bytes,
    *,
    preserve_active: bool = False,
) -> int:
    """Extract and replace a queued document's content and chunks idempotently.

    On a first ingest we activate the document so it becomes usable once ready.
    ``preserve_active=True`` (used by rescan) instead keeps whatever the admin
    last chose, so a disabled entry stays disabled after its content refreshes.
    """
    text = await asyncio.to_thread(extract_text, filename, data)
    chunks = chunk_text(text)
    if not chunks:
        raise ExtractionError("No useful text could be extracted from this file.")

    repo = KnowledgeRepository(session)
    await repo.delete_chunks(document.id)
    document.content = text
    for chunk in chunks:
        embedding = await embed_passage(chunk)
        await repo.add_chunk(
            profile_id=document.profile_id,
            document_id=document.id,
            content=chunk,
            embedding=embedding,
            meta={"title": document.title or filename},
        )
    document.processing_status = "ready"
    document.processing_error = None
    if not preserve_active:
        document.is_active = True
    document.processed_at = datetime.now(timezone.utc)
    await session.flush()
    return len(chunks)
